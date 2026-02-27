"""DOCX cover letter builder.

Generates a professionally formatted cover letter as a DOCX file.
"""

import io
import logging
from datetime import datetime
from typing import Dict

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def build_cover_letter_docx(cover_letter_text: str, user_profile: Dict = None) -> bytes:
    """Generate a professional DOCX cover letter.

    Args:
        cover_letter_text: Plain text cover letter content from Gemini.
        user_profile: Optional user profile for header formatting.

    Returns:
        DOCX file content as bytes.
    """
    doc = Document()
    _setup_styles(doc)

    # Parse the cover letter text into paragraphs and format
    paragraphs = cover_letter_text.strip().split("\n")

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            doc.add_paragraph("")  # Empty line for spacing
            continue

        p = doc.add_paragraph(para_text)

        # Detect and style specific parts
        if _is_date_line(para_text):
            p.paragraph_format.space_after = Pt(12)
        elif _is_greeting(para_text):
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif _is_closing(para_text):
            p.paragraph_format.space_before = Pt(12)
        else:
            p.paragraph_format.space_after = Pt(6)
            # Justify body paragraphs
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── Internal Helpers ───────────────────────────────────────────────────────────

def _setup_styles(doc: Document):
    """Set up default document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _is_date_line(text: str) -> bool:
    """Check if a line looks like a date."""
    import re
    return bool(re.match(
        r"^(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}",
        text
    ))


def _is_greeting(text: str) -> bool:
    """Check if a line is a greeting/salutation."""
    lower = text.lower()
    return lower.startswith(("dear ", "to whom", "hi ", "hello "))


def _is_closing(text: str) -> bool:
    """Check if a line is a closing."""
    lower = text.lower().rstrip(",.")
    return lower in (
        "sincerely", "best regards", "regards", "yours truly",
        "kind regards", "thank you", "respectfully", "warm regards",
    )
