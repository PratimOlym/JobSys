"""DOCX resume builder.

Generates a professionally formatted, ATS-optimized resume as a DOCX file
using python-docx and the structured data from Gemini.
"""

import io
import logging
from typing import Dict

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


def build_resume_docx(resume_data: Dict) -> bytes:
    """Generate a professional DOCX resume from structured data.

    Args:
        resume_data: Dict with sections (contact_info, professional_summary,
                     skills, experience, education, certifications).

    Returns:
        DOCX file content as bytes.
    """
    doc = Document()
    _setup_styles(doc)

    # ── Contact Header ─────────────────────────────────────────────────────
    contact = resume_data.get("contact_info", {})
    name = contact.get("name", "")
    if name:
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(name.upper())
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Contact details line
    contact_parts = []
    for field in ["email", "phone", "location", "linkedin"]:
        val = contact.get(field, "")
        if val:
            contact_parts.append(val)

    if contact_parts:
        contact_line = doc.add_paragraph()
        contact_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact_line.add_run(" | ".join(contact_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    _add_divider(doc)

    # ── Professional Summary ──────────────────────────────────────────────
    summary = resume_data.get("professional_summary", "")
    if summary:
        _add_section_heading(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph(summary)
        p.style.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    # ── Skills ────────────────────────────────────────────────────────────
    skills = resume_data.get("skills", [])
    if skills:
        _add_section_heading(doc, "SKILLS")
        # Display skills in a comma-separated format
        skills_text = " • ".join(skills)
        p = doc.add_paragraph(skills_text)
        p.style.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    # ── Experience ────────────────────────────────────────────────────────
    experience = resume_data.get("experience", [])
    if experience:
        _add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
        for exp in experience:
            # Job title and company line
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(exp.get("title", ""))
            title_run.bold = True
            title_run.font.size = Pt(11)

            company = exp.get("company", "")
            duration = exp.get("duration", "")
            if company or duration:
                meta = f" — {company}" if company else ""
                if duration:
                    meta += f" | {duration}"
                meta_run = title_para.add_run(meta)
                meta_run.font.size = Pt(10)
                meta_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            title_para.paragraph_format.space_after = Pt(2)
            title_para.paragraph_format.space_before = Pt(6)

            # Bullet points
            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                for run in bp.runs:
                    run.font.size = Pt(10)

    # ── Education ─────────────────────────────────────────────────────────
    education = resume_data.get("education", [])
    if education:
        _add_section_heading(doc, "EDUCATION")
        for edu in education:
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(edu.get("degree", ""))
            degree_run.bold = True
            degree_run.font.size = Pt(10)

            institution = edu.get("institution", "")
            year = edu.get("year", "")
            if institution or year:
                details = f" — {institution}" if institution else ""
                if year:
                    details += f" ({year})"
                details_run = edu_para.add_run(details)
                details_run.font.size = Pt(10)
                details_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Certifications ────────────────────────────────────────────────────
    certifications = resume_data.get("certifications", [])
    if certifications:
        _add_section_heading(doc, "CERTIFICATIONS")
        for cert in certifications:
            p = doc.add_paragraph(f"• {cert}")
            for run in p.runs:
                run.font.size = Pt(10)

    # ── Additional Sections ───────────────────────────────────────────────
    additional = resume_data.get("additional_sections", {})
    if additional and isinstance(additional, dict):
        for section_name, content in additional.items():
            _add_section_heading(doc, section_name.upper())
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(f"• {item}")
            elif isinstance(content, str):
                doc.add_paragraph(content)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Write to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── Internal Helpers ───────────────────────────────────────────────────────────

def _setup_styles(doc: Document):
    """Set up default document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_section_heading(doc: Document, text: str):
    """Add a styled section heading with underline."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Add a thin border below the heading
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "1A1A2E",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_divider(doc: Document):
    """Add a thin horizontal divider."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "CCCCCC",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
